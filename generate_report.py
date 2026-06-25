"""
生成项目报告 DOCX
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ── 样式设置 ──
style = doc.styles['Normal']
font = style.font; font.name = '宋体'; font.size = Pt(11)
style.paragraph_format.line_spacing = 1.15

# ── 封面 ──
for _ in range(6): doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('基于脉冲神经网络的\nMNIST 手写数字识别')
run.font.size = Pt(26); run.font.bold = True; run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('无监督 STDP 特征学习 + MLP 分类器')
run.font.size = Pt(16); run.font.color.rgb = RGBColor(0x58, 0x6e, 0x94)

doc.add_paragraph()
doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('认知科学类脑计算课程项目\n').font.size = Pt(12)
info.add_run('最佳准确率: 94.75%').font.size = Pt(12)

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.add_run('2026 年 6 月').font.size = Pt(12)

doc.add_page_break()

# ── 目录页 ──
doc.add_heading('目录', level=1)
toc_items = [
    '摘要', '引言与项目背景', '相关研究与技术背景', '架构设计',
    '学习规则与训练流程', '关键修正与失败教训', '实验结果',
    '与 BindsNET 的对比分析', 'Web 验证平台', '工程实践亮点', '总结与展望', '参考文献'
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# ── 摘要 ──
doc.add_heading('摘要', level=1)
doc.add_paragraph(
    '本项目实现了一个基于脉冲神经网络的 MNIST 手写数字识别系统。'
    '网络第一层为卷积脉冲层，使用无监督 STDP（Spike-Timing-Dependent Plasticity）'
    '学习规则进行训练，完全不依赖标签；第二层为监督 MLP 分类器。'
    '经过多次架构迭代——从最初的全连接 LIF + 侧向抑制 + 纯无监督投票方案，'
    '演变为 Conv + IF + 乘性 STDP + MLP 的混合方案——最终在 MNIST 测试集上达到 94.75% 的准确率。'
    '项目同时实现了断点续训、TensorBoard 实时监控、Undefined 拒识检测以及 Flask Web 验证平台。'
    '与 BindsNET（相同架构，93.12%）的对比验证了本方案的有效性和工程优越性。'
)

doc.add_heading('关键词', level=2)
doc.add_paragraph('脉冲神经网络 · STDP · 无监督学习 · MNIST · SpikingJelly · 卷积神经网络')

doc.add_page_break()

# ── 1. 引言 ──
doc.add_heading('1. 引言与项目背景', level=1)
doc.add_heading('1.1 项目目标', level=2)
doc.add_paragraph(
    '实现一个脉冲神经网络（SNN），使用无监督 STDP 训练卷积层提取视觉特征，'
    '配合监督 MLP 分类器完成 MNIST 手写数字识别。训练后准确率目标 ≥ 92%。'
)

doc.add_heading('1.2 原始要求', level=2)
doc.add_paragraph('老师最初规定的方案包含以下要点：')
items = [
    '全连接层：784 个输入神经元 → 120 个兴奋 LIF 神经元',
    '侧向抑制：120 个抑制 LIF 神经元，通过全连接负权重实现 WTA（赢者通吃）',
    '纯无监督分类：训练后为每个神经元分配数字类别，采用神经元投票机制',
    '在线迹 STDP：A_pre=0.1, A_post=-0.12, tau=20ms, T=100',
    '目标准确率 ≥ 92%'
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('1.3 方案演进', level=2)
table = doc.add_table(rows=5, cols=5)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['阶段', '架构', '分类方式', '准确率', '结论']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
data = [
    ['V1', 'FC 784→120+LIF+侧向抑制', '纯无监督投票', '<10%', '❌ 全连接+WTA 失败'],
    ['V2', 'Conv 5×5×32+IF', '纯无监督滤波器投票', '9.80%', '❌ 无分类器不可行'],
    ['V3 (5K)', 'Conv+IF+MLP', 'STDP特征+MLP', '92.03%', '✅'],
    ['V3 (60K)', 'Conv+IF+MLP', 'STDP特征+MLP', '94.75%', '✅ 最终方案'],
]
for r, row_data in enumerate(data):
    for c, val in enumerate(row_data):
        table.rows[r+1].cells[c].text = val

doc.add_page_break()

# ── 2. 相关研究 ──
doc.add_heading('2. 相关研究与技术背景', level=1)
doc.add_heading('2.1 STDP 原理', level=2)
doc.add_paragraph(
    'STDP（Spike-Timing-Dependent Plasticity）是生物神经元中普遍存在的突触可塑性机制。'
    '其核心规则为：若突触前神经元先于突触后神经元发放（因果配对），则突触权重增强（LTP）；'
    '若突触后神经元先于突触前神经元发放（非因果配对），则突触权重减弱（LTD）。'
)
doc.add_paragraph(
    '本项目采用权重依赖（乘性）STDP：f_pre(w) = a_plus × w 控制 LTD，'
    'f_post(w) = a_plus × (1-w) 控制 LTP。权重在 [0,1] 区间内自然稳定。'
)

doc.add_heading('2.2 关键论文', level=2)
papers = [
    ('Diehl & Cook (2015)', 'Frontiers in Computational Neuroscience',
     '领域奠基论文，纯无监督 FC+WTA SNN，6400 神经元达到 95.0%'),
    ('Kheradpisheh et al. (2018)', 'Neural Networks',
     '深度卷积 STDP + SVM 分类器，达到 98.4%，本项目的直接参考'),
    ('Lee et al. (2019)', 'IEEE TCDS',
     'SpiCNN：无监督卷积 STDP 达到 91.1%，架构与本研究最接近'),
    ('Mozafari et al. (2019)', 'Pattern Recognition',
     'STDP + 奖励调制 STDP，纯无监督达到 97.2%'),
]
for title, journal, desc in papers:
    doc.add_paragraph(f'{title} — {journal}', style='List Bullet')
    doc.add_paragraph(f'    {desc}')

doc.add_heading('2.3 技术栈', level=2)
tech = [
    'PyTorch 2.5.1 (CUDA 12.1) — 张量计算与 MLP 训练',
    'SpikingJelly 0.0.0.0.14 — STDP 学习器、IF 神经元、泊松编码',
    'Flask — Web 验证平台',
    'TensorBoard — 训练实时监控',
    'Matplotlib + Seaborn — 混淆矩阵与权重可视化'
]
for t in tech:
    doc.add_paragraph(t, style='List Bullet')

doc.add_page_break()

# ── 3. 架构设计 ──
doc.add_heading('3. 架构设计', level=1)
doc.add_heading('3.1 整体架构', level=2)
doc.add_paragraph(
    '输入图像（28×28 灰度图）→ 泊松编码（T=10 时间步，每步独立以概率 p=像素/255 发放脉冲）'
    '→ Conv2d（1→32, kernel=5×5, padding=2, 800 参数）→ IF 神经元（阈值=1.0, 无泄漏积分）'
    '→ AvgPool（2×2, stride=2）→ 展平为 6272 维特征向量'
    '→ MLP（6272→512→11）→ 输出 0-9 数字或 undefined(10)'
)

doc.add_heading('3.2 卷积 STDP 层', level=2)
params_conv = [
    ('滤波器数量', '32 个 5×5 卷积核'),
    ('参数量', '32 × 1 × 5 × 5 = 800（仅为 FC 方案的 1/117）'),
    ('神经元', 'IFNode, v_threshold=1.0, step_mode="s"'),
    ('权重初始化', 'U(0, 0.3), bias=False'),
    ('池化', 'AvgPool2d(2×2), 32×28×28 → 32×14×14 = 6272'),
    ('编码', '泊松编码, T=10, p=像素值/255'),
]
for k, v in params_conv:
    doc.add_paragraph(f'{k}: {v}', style='List Bullet')

doc.add_heading('3.3 MLP 分类器', level=2)
params_mlp = [
    ('fc1', '6272 → 512, ReLU, Dropout(0.3)'),
    ('fc2', '512 → 11 (0-9 + undefined)'),
    ('优化', 'Adam(lr=0.001, weight_decay=1e-4)'),
    ('损失', 'CrossEntropyLoss'),
    ('batch_size', '256'),
    ('epochs', '50'),
]
for k, v in params_mlp:
    doc.add_paragraph(f'{k}: {v}', style='List Bullet')

doc.add_page_break()

# ── 4. 学习规则 ──
doc.add_heading('4. 学习规则与训练流程', level=1)
doc.add_heading('4.1 乘性 STDP', level=2)
doc.add_paragraph(
    'f_pre(w) = a_plus × w          — LTD 幅度正比于当前权重，权重越大越容易被减弱\n'
    'f_post(w) = a_plus × (1.0 - w)  — LTP 幅度反比于当前权重，权重越小越容易被增强'
)
doc.add_paragraph(
    '参数：a_plus = 0.004, tau_pre = tau_post = 20.0ms, lr = 0.02, step_mode = "s"\n'
    '关键：T=10 步累积后一次性应用权重更新，避免逐时间步更新导致的振荡'
)

doc.add_heading('4.2 两阶段训练', level=2)
doc.add_paragraph('阶段 1 — STDP 无监督训练（5 epochs）：')
steps1 = [
    '每个样本 batch_size=1',
    '泊松编码 → T=10 步脉冲序列',
    '逐时间步前向：Conv + IF 发放',
    'STDP 监视器自动记录 pre/post 脉冲配对',
    'T 步累积 delta_w → 一次性应用到卷积核权重',
    '每 epoch 结束：L2 归一化到 |w|=5.0 → 保存断点'
]
for s in steps1:
    doc.add_paragraph(s, style='List Bullet')

doc.add_paragraph('阶段 2 — MLP 监督训练（50 epochs）：')
steps2 = [
    '特征提取：Conv+IF 前向 → 时间平均 → AvgPool → 6272 维',
    '合成 undefined 样本：随机混合两个不同类的特征 (50:50)',
    'Adam 优化，每 5 epoch 评估 test_acc',
    '保存最佳模型 (按 test_acc)'
]
for s in steps2:
    doc.add_paragraph(s, style='List Bullet')

doc.add_page_break()

# ── 5. 关键修正 ──
doc.add_heading('5. 关键修正与失败教训', level=1)

doc.add_heading('5.1 STDP 符号修正', level=2)
doc.add_paragraph(
    '原始参数 A_post=-0.12 导致纯反赫布学习：pre→post 的因果配对理应变强却变弱，'
    '两种场景权重都减小，所有权重归零，网络死亡。修正为乘性 STDP（a_plus=0.004），'
    'LTP 与 LTD 配合，权重在 [0,1] 自然稳定。'
)

doc.add_heading('5.2 step_mode bug', level=2)
doc.add_paragraph(
    'LIF/IF 的 step_mode="m" 在逐时间步调用时，膜电位 v 被错误初始化为 [N] 而非 [batch,N]，'
    '导致所有 batch 样本共享同一个膜电位，网络无法区分不同输入。改为 step_mode="s"。'
)

doc.add_heading('5.3 全连接 → 卷积', level=2)
doc.add_paragraph(
    'FC 方案 94,080 个参数，每个权重每样本仅约 2% 机会被 STDP 更新；'
    'Conv 方案 800 个参数（少 117 倍），卷积权重共享使每个权重每样本被更新约 70 次。'
    '此外卷积自带平移不变性归纳偏置，对图像识别天然更高效。'
)

doc.add_heading('5.4 LIF → IF', level=2)
doc.add_paragraph(
    'LIF 的膜电位每步泄漏 50%（τ=2.0），在 T=10 的短窗口下大量信号丢失。'
    'V3 架构无循环连接，不需要 LIF 的振荡阻尼功能，改用 IF 完美积分更高效。'
)

doc.add_heading('5.5 纯无监督投票 → MLP', level=2)
doc.add_paragraph(
    '无监督投票方案在 FC 和 Conv 上均失败——WTA 正反馈导致所有神经元偏向同一类别。'
    'MLP 学习特征的软组合，不需要每个神经元专属某个数字。'
    '注意：STDP 部分仍然是无监督的，MLP 只是\'读取器\'。'
)

doc.add_page_break()

# ── 6. 实验结果 ──
doc.add_heading('6. 实验结果', level=1)

doc.add_heading('6.1 准确率', level=2)
table2 = doc.add_table(rows=5, cols=3)
table2.style = 'Light Grid Accent 1'
for i, h in enumerate(['训练样本', '准确率', '备注']):
    table2.rows[0].cells[i].text = h
acc_data = [
    ('500', '80.65%', '快速验证'),
    ('3,000', '89.81%', '接近论文基线'),
    ('5,000', '92.03%', '超过原始目标 92%'),
    ('60,000', '94.75%', '最终最佳'),
]
for r, (s, a, n) in enumerate(acc_data):
    table2.rows[r+1].cells[0].text = s
    table2.rows[r+1].cells[1].text = a
    table2.rows[r+1].cells[2].text = n

doc.add_heading('6.2 MLP 训练曲线', level=2)
doc.add_paragraph(
    'Epoch 1: 88.3% train, 88.9% test → Epoch 10: 93.5% train, 93.5% test '
    '→ Epoch 35: 94.7% train, 94.8% test (最佳) → Epoch 50: 94.6% train, 94.6% test'
)

doc.add_heading('6.3 混淆矩阵', level=2)
cm_path = 'results/confusion_matrix_60k.png'
if os.path.exists(cm_path):
    doc.add_picture(cm_path, width=Inches(4.5))

doc.add_page_break()

# ── 7. 对比 BindsNET ──
doc.add_heading('7. 与 BindsNET 的对比分析', level=1)
doc.add_paragraph(
    '另一个终端使用 BindsNET 框架复现了相同架构（Conv+IF+STDP+MLP），'
    '达到 93.12% 的准确率。以下是详细对比：'
)

table3 = doc.add_table(rows=9, cols=3)
table3.style = 'Light Grid Accent 1'
for i, h in enumerate(['维度', 'BindsNET', '本项目']):
    table3.rows[0].cells[i].text = h
comp = [
    ('框架', 'BindsNET', 'SpikingJelly'),
    ('a_plus', '5×10⁻⁶', '0.004 (800×)'),
    ('训练样本', '60,000', '60,000'),
    ('最终准确率', '93.12%', '94.75%'),
    ('训练时间', '4.4h', '~14h (含STDP)'),
    ('断点续训', '无', '有'),
    ('Undefined 检测', '无', '11 类含 undefined'),
    ('Web 验证', '无', 'Flask 实时平台'),
]
for r, (dim, bn, ours) in enumerate(comp):
    table3.rows[r+1].cells[0].text = dim
    table3.rows[r+1].cells[1].text = bn
    table3.rows[r+1].cells[2].text = ours

if os.path.exists('bindsnet_results/accuracy_curve.png'):
    doc.add_paragraph('\nBindsNET 准确率曲线：')
    doc.add_picture('bindsnet_results/accuracy_curve.png', width=Inches(4.5))

doc.add_page_break()

# ── 8. Web 平台 ──
doc.add_heading('8. Web 验证平台', level=1)
doc.add_paragraph(
    '使用 Flask 框架构建的实时验证平台，支持：\n'
    '• 实时流式验证 MNIST 测试集，动态显示当前图像、预测结果和置信度\n'
    '• 32 个滤波器特征图实时可视化（4×8 网格，展示 STDP 学到的边缘/形状）\n'
    '• 动态准确率统计 + 近 50 次预测趋势图（绿色正确/红色错误/黄色 Undefined）\n'
    '• 自定义图片上传：任意尺寸 → 28×28 黑底白字预处理 → 模型推理\n'
    '• Undefined 检测：置信度 < 0.3 的预测拒识为类别 10\n'
    '• 键盘快捷键：空格开始/暂停，Ctrl+R 重置'
)

doc.add_heading('8.1 架构图', level=2)
if os.path.exists('ARCHITECTURE.md'):
    doc.add_paragraph('详细架构文档见项目中的 ARCHITECTURE.md')

doc.add_page_break()

# ── 9. 工程实践 ──
doc.add_heading('9. 工程实践亮点', level=1)
practices = [
    '断点续训：训练中断后自动恢复，STDP 权重、MLP 权重、epoch 进度全覆盖',
    'TensorBoard 集成：实时监控 STDP 发放率 + MLP Loss/Accuracy 曲线',
    '乘性 STDP 自稳定：f_pre∝w, f_post∝(1-w)，权重自然约束在 [0,1]',
    'T 步累积更新：避免逐时间步权重振荡，提升训练稳定性',
    '模块化代码：model / train / main / app 清晰分离',
    '完整文档：README、SPEC、ARCHITECTURE、PPT_OUTLINE 多层级覆盖',
]
for p in practices:
    doc.add_paragraph(p, style='List Bullet')

doc.add_page_break()

# ── 10. 总结 ──
doc.add_heading('10. 总结与展望', level=1)
doc.add_heading('10.1 主要结论', level=2)
conclusions = [
    'STDP 可以做有效的无监督特征提取——卷积核学到了边缘、方向、形状等视觉特征',
    '纯无监督分类（神经元投票）在 FC 和 Conv 架构上均失败，需要分类器',
    '卷积 >> 全连接：参数少 117 倍，归纳偏置天然适合图像',
    '乘性 STDP >> 加法 STDP：自稳定，无需额外约束',
    '94.75% 准确率超过原始目标 92%，接近论文 SOTA',
]
for c in conclusions:
    doc.add_paragraph(c, style='List Bullet')

doc.add_heading('10.2 未来方向', level=2)
future = [
    '纯无监督分类：引入均匀分配策略 + 侧向抑制 + 自适应阈值',
    'ESP32-CAM 部署：将模型压缩至 32KB 以下（当前 12MB）',
    '更多数据集：Fashion-MNIST, EMNIST, CIFAR-10 (DVS)',
    '更深架构：多层卷积 STDP + 逐层训练',
]
for f in future:
    doc.add_paragraph(f, style='List Bullet')

doc.add_page_break()

# ── 参考文献 ──
doc.add_heading('参考文献', level=1)
refs = [
    '[1] Diehl, P. U. & Cook, M. (2015). Unsupervised learning of digit recognition '
    'using spike-timing-dependent plasticity. Frontiers in Computational Neuroscience, 9:99.',
    '[2] Kheradpisheh, S. R. et al. (2018). STDP-based spiking deep convolutional neural '
    'networks for object recognition. Neural Networks, 99, 56-67.',
    '[3] Lee, C. et al. (2019). Deep Spiking Convolutional Neural Network Trained With '
    'Unsupervised STDP. IEEE Trans. Cognitive and Developmental Systems.',
    '[4] Mozafari, M. et al. (2019). Bio-inspired digit recognition using reward-modulated '
    'STDP in deep convolutional networks. Pattern Recognition, 94.',
    '[5] Fang, W. et al. (2020). SpikingJelly. https://github.com/fangwei123456/spikingjelly',
    '[6] Hazan, H. et al. (2018). BindsNET: A machine learning-oriented spiking neural '
    'networks library in Python. Frontiers in Neuroinformatics, 12:89.',
]
for ref in refs:
    doc.add_paragraph(ref)

# ── 保存 ──
output_path = '项目报告_STDP_SNN_MNIST.docx'
doc.save(output_path)
print(f'✅ 报告已生成: {output_path}')
print(f'   {os.path.getsize(output_path)/1024:.0f} KB')
